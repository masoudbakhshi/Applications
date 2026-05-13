"""
report_generator.py — Word (.docx) report generation
Author: Masoud Bakhshi
"""

from __future__ import annotations
import io
import datetime
from typing import Dict, Any, Optional, List

import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


AUTHOR = "Masoud Bakhshi"
APP_NAME = "CurrentControllerDesigner"
VERSION = "1.0"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def _para(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def _table_2col(doc: Document, rows: List[tuple]):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Value"
    for k, v in rows:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)
    return table


def _add_figure(doc: Document, fig: plt.Figure, caption: str = ""):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    doc.add_picture(buf, width=Inches(5.5))
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
    plt.close(fig)


# ---------------------------------------------------------------------------
# Main report builder
# ---------------------------------------------------------------------------

def generate_report(
    machine_type: str,
    params: Dict[str, Any],
    design_results: Dict[str, Any],
    stability_reports: Dict[str, Any],
    figures: Dict[str, plt.Figure],
    discrete_results: Dict[str, Any],
    c_snippets: Dict[str, str],
) -> bytes:
    """
    Build and return a Word document as bytes.

    Parameters
    ----------
    machine_type    : 'EESM' | 'IPMSM' | 'IM'
    params          : dict of user input parameters
    design_results  : dict with PIController objects per channel
    stability_reports : dict with StabilityReport objects per channel
    figures         : dict mapping label → matplotlib Figure
    discrete_results: dict with DiscretePIResult objects per channel
    c_snippets      : dict with embedded C code snippets per channel
    """
    doc = Document()

    # ---- Cover ----
    doc.add_heading(f"{APP_NAME} — Controller Design Report", 0)
    doc.add_heading(f"Machine Type: {machine_type}", 2)
    _para(doc, f"Author: {AUTHOR}")
    _para(doc, f"Application: {APP_NAME} v{VERSION}")
    _para(doc, f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()
    _para(doc,
          "DISCLAIMER: This report is generated automatically based on the "
          "parameters entered by the user. All results must be validated "
          "against measured machine data and hardware-in-the-loop testing "
          "before deployment in a safety-critical automotive system.",
          italic=True)
    doc.add_page_break()

    # ---- Table of Contents placeholder ----
    _heading(doc, "1.  Introduction", 1)
    _para(doc,
          f"This report documents the current-loop, torque-loop, and "
          f"optional speed-loop controller design for a {machine_type} "
          f"propulsion drive. The design was performed using the "
          f"{APP_NAME} tool authored by {AUTHOR}.")

    # ---- Machine Parameters ----
    _heading(doc, "2.  Machine and Inverter Parameters", 1)
    _para(doc,
          "The following parameters were entered by the user. Parameters "
          "marked with (*) use default example values and must be replaced "
          "with measured or validated data.")
    param_rows = [(k, _fmt(v)) for k, v in params.items()]
    _table_2col(doc, param_rows)
    doc.add_paragraph()

    # ---- Mathematical Model ----
    _heading(doc, "3.  Mathematical Model", 1)
    _add_model_section(doc, machine_type)

    # ---- Controller Design ----
    _heading(doc, "4.  Continuous-Time Controller Design", 1)
    for ch_name, pi in design_results.items():
        _heading(doc, f"4.{list(design_results.keys()).index(ch_name)+1}  {ch_name} Channel", 2)
        if pi is None:
            _para(doc, "Not applicable for this machine type.")
            continue
        rows = [
            ("Design Method",  pi.get("method", "—")),
            ("Kp",             f"{pi['Kp']:.6g}"),
            ("Ki",             f"{pi['Ki']:.6g}"),
            ("Ti (= Kp/Ki)",   f"{pi['Ti']:.6g} s"),
            ("Target BW",      f"{pi.get('omega_bw', '—')} rad/s"),
        ]
        _table_2col(doc, rows)
        doc.add_paragraph()

    # ---- Discrete-Time Design ----
    _heading(doc, "5.  Discrete-Time Controller Design", 1)
    for ch_name, dr in discrete_results.items():
        _heading(doc, f"5.{list(discrete_results.keys()).index(ch_name)+1}  {ch_name} Channel", 2)
        if dr is None:
            _para(doc, "Not applicable.")
            continue
        rows = [
            ("Discretization Method", dr.method),
            ("Sampling Period Ts",    f"{dr.Ts*1e6:.2f} µs"),
            ("Kp (same as CT)",       f"{dr.Kp:.6g}"),
            ("Ki continuous",         f"{dr.Ki_c:.6g}"),
            ("b0",                    f"{dr.b0:.6g}"),
            ("b1",                    f"{dr.b1:.6g}"),
            ("Ki discrete",           f"{dr.Ki_d:.6g}"),
        ]
        _table_2col(doc, rows)
        _para(doc, "Difference equation (position form):", bold=True)
        _para(doc,
              f"  u[k] = u[k-1] + {dr.b0:.6g}·e[k] + {dr.b1:.6g}·e[k-1]",
              italic=True)
        doc.add_paragraph()

    # ---- Stability Analysis ----
    _heading(doc, "6.  Stability and Robustness Analysis", 1)
    for ch_name, sr in stability_reports.items():
        _heading(doc, f"6.{list(stability_reports.keys()).index(ch_name)+1}  {ch_name}", 2)
        if sr is None:
            _para(doc, "Not applicable.")
            continue
        rows = [
            ("Phase Margin",           f"{sr.phase_margin_deg:.1f}°"),
            ("Gain Margin",            f"{sr.gain_margin_db:.1f} dB"),
            ("Gain Crossover Freq",    f"{sr.gain_crossover_freq:.1f} rad/s"),
            ("Closed-Loop BW (−3 dB)", f"{sr.bandwidth_cl:.1f} rad/s"),
            ("Sensitivity Peak Ms",    f"{sr.Ms:.3f}"),
            ("Rise Time",              f"{sr.rise_time_ms:.2f} ms"),
            ("Settling Time (±2%)",    f"{sr.settling_time_ms:.2f} ms"),
            ("Overshoot",              f"{sr.overshoot_pct:.1f}%"),
            ("Stable (CT)",            str(sr.continuous_stable)),
            ("Stable (DT)",            str(sr.discrete_stable)),
        ]
        _table_2col(doc, rows)
        if sr.warnings:
            _para(doc, "Warnings:", bold=True)
            for w in sr.warnings:
                _para(doc, f"  • {w}")
        doc.add_paragraph()

    # ---- Figures ----
    _heading(doc, "7.  Plots", 1)
    for fig_label, fig in figures.items():
        if fig is not None:
            _add_figure(doc, fig, caption=fig_label)
            doc.add_paragraph()

    # ---- Embedded C Code ----
    _heading(doc, "8.  Embedded Implementation (C Code Snippets)", 1)
    _para(doc,
          "The following ANSI-C snippets implement the discrete PI "
          "controllers suitable for an automotive ECU. Anti-windup "
          "(back-calculation method) is included.")
    for ch_name, snippet in c_snippets.items():
        _heading(doc, ch_name, 2)
        p = doc.add_paragraph()
        run = p.add_run(snippet)
        run.font.name = "Courier New"
        run.font.size = Pt(8)

    # ---- Practical Notes ----
    _heading(doc, "9.  Practical Implementation Notes", 1)
    _add_practical_notes(doc, machine_type)

    # ---- Assumptions & Limitations ----
    _heading(doc, "10. Assumptions and Limitations", 1)
    _add_limitations(doc)

    # ---- Footer ----
    doc.add_page_break()
    _para(doc, f"Report generated by {APP_NAME} — Author: {AUTHOR}", italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Sub-sections
# ---------------------------------------------------------------------------

def _add_model_section(doc: Document, machine_type: str):
    if machine_type == "EESM":
        _para(doc, "dq-axis voltage equations (motor convention):", bold=True)
        eqs = [
            "v_d = R_s·i_d + L_d·(di_d/dt) − ω_e·L_q·i_q",
            "v_q = R_s·i_q + L_q·(di_q/dt) + ω_e·(L_d·i_d + ψ_f(i_f))",
            "v_f = R_f·i_f + L_f·(di_f/dt)  (field winding)",
        ]
        for eq in eqs:
            _para(doc, f"  {eq}", italic=True)
        _para(doc, "Electromagnetic torque:", bold=True)
        _para(doc,
              "  T_e = (3/2)·p·[ψ_f(i_f)·i_q + (L_d − L_q)·i_d·i_q]",
              italic=True)
        _para(doc,
              "The field flux linkage ψ_f(i_f) is either modelled as "
              "ψ_f = L_m·i_f (linear) or provided as a lookup table.")

    elif machine_type == "IPMSM":
        _para(doc, "dq-axis voltage equations:", bold=True)
        eqs = [
            "v_d = R_s·i_d + L_d·(di_d/dt) − ω_e·L_q·i_q",
            "v_q = R_s·i_q + L_q·(di_q/dt) + ω_e·(L_d·i_d + ψ_pm)",
        ]
        for eq in eqs:
            _para(doc, f"  {eq}", italic=True)
        _para(doc, "Electromagnetic torque:", bold=True)
        _para(doc,
              "  T_e = (3/2)·p·[ψ_pm·i_q + (L_d − L_q)·i_d·i_q]",
              italic=True)

    elif machine_type == "IM":
        _para(doc, "Rotor-flux-oriented voltage equations:", bold=True)
        eqs = [
            "v_sd = R_s·i_sd + σ·L_s·(di_sd/dt) − ω_s·σ·L_s·i_sq",
            "v_sq = R_s·i_sq + σ·L_s·(di_sq/dt) + ω_s·(σ·L_s·i_sd + ψ_r/L_s)",
            "dψ_r/dt = (L_m/T_r)·i_sd − (1/T_r)·ψ_r",
        ]
        for eq in eqs:
            _para(doc, f"  {eq}", italic=True)
        _para(doc, "Electromagnetic torque:", bold=True)
        _para(doc,
              "  T_e = (3/2)·p·(L_m/L_r)·ψ_r·i_sq",
              italic=True)
        _para(doc, "Slip frequency:", bold=True)
        _para(doc,
              "  ω_slip = (L_m/T_r)·i_sq / ψ_r",
              italic=True)

    _para(doc, "Delay model:", bold=True)
    _para(doc,
          "  T_total = T_comp + T_pwm/2 + T_adc + T_filter\n"
          "  Modelled as a Padé approximant of order 2: e^{−sT_total}",
          italic=True)


def _add_practical_notes(doc: Document, machine_type: str):
    notes = [
        "1. Feedforward / decoupling: Cross-coupling terms (ω_e·L_q·i_q "
           "and ω_e·L_d·i_d) should be computed from measured speed and "
           "added to the PI output as feedforward before the PWM modulator.",
        "2. Anti-windup: The back-calculation method is used. The gain "
           "K_aw = 1/Kp should be tuned on the target platform.",
        "3. Voltage saturation: Apply current-limited MTPA/field-weakening "
           "algorithms to ensure the reference voltage never exceeds "
           "V_lim = Vdc / √3 (space-vector modulation limit).",
        "4. Sampling jitter: For symmetric PWM, sample at the PWM counter "
           "peak to minimise ripple aliasing.",
        "5. Dead-time: At high currents, dead-time compensation improves "
           "current quality. Model as ΔV_dt = V_dc·t_dt·f_pwm·sign(i).",
        "6. Temperature: R_s increases ~40% from 20°C to 120°C for copper. "
           "Adaptive gain scheduling or online resistance estimation is "
           "recommended for automotive applications.",
    ]
    if machine_type == "EESM":
        notes.append(
            "7. Field current dynamics: The rotor time constant T_f = L_f/R_f "
               "is typically much larger than the stator time constants. "
               "The field PI loop bandwidth must be designed accordingly."
        )
    elif machine_type == "IM":
        notes.append(
            "7. Slip estimation accuracy: The rotor time constant T_r = L_r/R_r "
               "is temperature-sensitive. Online T_r estimation (MRAS or "
               "Kalman filter) is recommended for production drives."
        )
    for note in notes:
        _para(doc, note)


def _add_limitations(doc: Document):
    items = [
        "• This tool performs linear, frequency-domain analysis only. "
          "Non-linearities (magnetic saturation, dead-time, PWM harmonics) "
          "are not fully captured.",
        "• The Padé delay approximation introduces phase error at high "
          "frequencies. Use order ≥ 2 for meaningful phase margin results.",
        "• Field-weakening operation requires additional outer-loop logic "
          "not covered in this basic design.",
        "• The discrete-time stability analysis uses approximate z-domain "
          "models. Exact validation should be performed via simulation.",
        "• All results assume balanced three-phase operation.",
        "• Mutual inductances between stator and field winding are "
          "approximated; use FEA-derived parameters for precision.",
    ]
    for item in items:
        _para(doc, item)


def _fmt(v) -> str:
    if isinstance(v, float):
        return f"{v:.6g}"
    return str(v)
